"""生成 demo 题：problems/demo-cumcm/

题目：小型制造厂"7 天 × 3 班次多机器排班 + 产量预测"
- Q1：排班优化（整数规划，ortools CP-SAT）—— 覆盖优化路径
- Q2：产量预测（线性回归 + 拟合评估）—— 覆盖预测路径
数据带已知最优结构，便于核对。

用法：$PYTHON_EXE make_demo.py [project_root]
"""
from __future__ import annotations

import pathlib
import sys

import numpy as np

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))

import machine as M

M.ensure_utf8_stdout()

# 本文件位于 .claude/skills/math-modeling-workflow/scripts/ → 项目根在 parents[4]
PROJECT_ROOT = pathlib.Path(__file__).resolve().parents[4]
DEMO_DIR = PROJECT_ROOT / "problems" / "demo-cumcm"

N_PRODUCTS, N_DAYS = 4, 7
RATES = np.array([  # 生产率（件/小时）：机器 M1 M2 M3
    [4, 3, 2],
    [3, 4, 2],
    [2, 3, 4],
    [3, 2, 3],
], dtype=int)
COST = np.array([8, 10, 12, 9])  # 单位生产成本（元/件）
HOURS_PER_SHIFT = 8
SHIFTS_PER_DAY = 3


def gen_csv(demo_dir: pathlib.Path) -> None:
    rng = np.random.default_rng(42)
    # 附件1：需求表（4 产品 × 7 天）
    demand = rng.integers(30, 80, size=(N_PRODUCTS, N_DAYS))
    lines = ["type,day,demand"]
    for p in range(N_PRODUCTS):
        for d in range(N_DAYS):
            lines.append(f"P{p+1},{d+1},{demand[p, d]}")
    (demo_dir / "attachment1_demand.csv").write_text("\n".join(lines) + "\n", encoding="utf-8")

    # 附件2：机器与产品参数
    lines = ["product,machine1_rate,machine2_rate,machine3_rate,unit_cost"]
    for p in range(N_PRODUCTS):
        lines.append(f"P{p+1},{RATES[p,0]},{RATES[p,1]},{RATES[p,2]},{COST[p]}")
    (demo_dir / "attachment2_params.csv").write_text("\n".join(lines) + "\n", encoding="utf-8")

    # 附件3：历史产量（Q2 预测）
    lines = ["day,output"]
    t = np.arange(1, 31)
    trend = 120 + 2.0 * t
    seasonal = 15 * np.sin(2 * np.pi * t / 7)
    hist = np.round(trend + seasonal + rng.normal(0, 4, size=t.size)).astype(int)
    for i, v in enumerate(hist, start=1):
        lines.append(f"{i},{v}")
    (demo_dir / "attachment3_history.csv").write_text("\n".join(lines) + "\n", encoding="utf-8")
    print("  已生成 attachment1_demand.csv / attachment2_params.csv / attachment3_history.csv")


def gen_docx(demo_dir: pathlib.Path) -> None:
    try:
        from docx import Document
    except ImportError:
        print("  python-docx 不可用，跳过 statement.docx（请手动补题目）")
        return
    doc = Document()
    doc.add_heading("数学建模全流程 Demo 题：小型制造厂排班与产量预测", level=0)
    doc.add_paragraph("（本题目由 make_demo.py 生成，用于端到端验证工作流，数据带已知结构。）")

    doc.add_heading("问题背景", level=1)
    doc.add_paragraph(
        "某小型制造厂有 3 台机器（M1、M2、M3），每天分早、中、晚 3 个班次，"
        "每班次每台机器可用 8 小时，连续运行 7 天。工厂生产 4 种产品（P1–P4），"
        "每种产品在不同机器上的生产效率不同，单位生产成本也不同。"
        "每天每种产品有固定的需求量，必须全部满足（不足可停产但按缺货高价惩罚）。")

    doc.add_heading("问题一：排班优化", level=1)
    doc.add_paragraph(
        "附件1给出了 4 种产品在 7 天内的每日需求量，附件2给出了各机器对各产品的生产效率"
        "（件/小时）以及单位生产成本（元/件）。请建立数学模型，安排各机器在各天各班次的生产计划，"
        "使 7 天总生产成本最小。要求给出每天每班次的排班表与总成本。")

    doc.add_heading("问题二：产量预测", level=1)
    doc.add_paragraph(
        "附件3给出了该厂过去 30 天的日产量（件）。请建立模型预测未来 5 天的产量，"
        "给出预测值，并用适当的指标评估模型的拟合效果。")

    doc.add_heading("数据说明", level=1)
    doc.add_paragraph("附件1（attachment1_demand.csv）：需求表，列 type, day, demand。")
    doc.add_paragraph("附件2（attachment2_params.csv）：产品-机器参数，列 product, machine1_rate, machine2_rate, machine3_rate, unit_cost。")
    doc.add_paragraph("附件3（attachment3_history.csv）：历史产量，列 day, output。")
    doc.add_paragraph("每台机器每班次可用小时数 = 8，每天 3 班次。")
    doc.save(str(demo_dir / "statement.docx"))
    print("  已生成 statement.docx")


def main() -> None:
    DEMO_DIR.mkdir(parents=True, exist_ok=True)
    print(f"==> 生成 demo 题到 {DEMO_DIR}")
    gen_csv(DEMO_DIR)
    gen_docx(DEMO_DIR)
    print("==> demo 题生成完成")


if __name__ == "__main__":
    main()
